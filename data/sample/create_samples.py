"""Generate synthetic sample contracts for testing.

Run this script once to create sample PDF and DOCX files in data/sample/.
"""

import os
from pathlib import Path

SAMPLE_DIR = Path(__file__).parent

MSA_TEXT = """MASTER SERVICES AGREEMENT

This Master Services Agreement ("Agreement") is entered into as of January 15, 2026 \
("Effective Date") by and between:

Acme Technology Solutions, Inc., a Delaware corporation ("Provider"), and
GlobalCorp Industries, LLC, a New York limited liability company ("Client").

ARTICLE 1: DEFINITIONS AND SCOPE OF SERVICES

1.1 Services. Provider shall provide Client with cloud-based software development and hosting \
services as described in the applicable Statement of Work ("SOW") executed by the parties.

1.2 Service Levels. Provider shall maintain system availability of 99.9% measured monthly, \
excluding scheduled maintenance windows. Scheduled maintenance shall not exceed 4 hours per \
month and shall occur between 2:00 AM and 6:00 AM Eastern Time with 48 hours advance notice.

ARTICLE 2: TERM AND TERMINATION

2.1 Term. This Agreement shall commence on the Effective Date and continue for an initial term \
of three (3) years ("Initial Term"), unless earlier terminated as provided herein.

2.2 Termination for Convenience. Either party may terminate this Agreement for convenience upon \
ninety (90) days prior written notice to the other party. Client shall pay for all services \
rendered through the effective date of termination plus a termination fee equal to three (3) \
months of fees.

2.3 Termination for Cause. Either party may terminate this Agreement immediately upon written \
notice if the other party: (a) materially breaches this Agreement and fails to cure such breach \
within thirty (30) days of receiving written notice; or (b) becomes insolvent, files for \
bankruptcy, or ceases operations.

2.4 Effect of Termination. Upon termination or expiration: (a) Provider shall return or destroy \
all Client Data within thirty (30) days; (b) Client shall pay all outstanding invoices within \
fifteen (15) days; (c) Sections 3, 5, 6, 7, and 9 shall survive termination.

ARTICLE 3: CONFIDENTIALITY

3.1 Confidential Information. Each party agrees to maintain the confidentiality of the other \
party's Confidential Information and not to disclose it to any third party without prior written \
consent, except as required by law.

3.2 Duration. Confidentiality obligations shall survive for five (5) years after termination of \
this Agreement, or indefinitely for trade secrets.

ARTICLE 4: DATA PROTECTION AND PRIVACY

4.1 Data Processing. Provider shall process Client personal data only as necessary to perform the \
Services and in accordance with applicable data protection laws including GDPR and CCPA.

4.2 Data Breach Notification. In the event of a data breach affecting Client personal data, \
Provider shall: (a) notify Client within forty-eight (48) hours of discovery; (b) provide a \
detailed incident report within five (5) business days; (c) cooperate with Client in notifying \
affected individuals and regulatory authorities.

4.3 Subprocessors. Provider shall not engage subprocessors to process Client personal data \
without Client's prior written consent. Provider shall remain liable for the acts and omissions \
of its subprocessors.

4.4 Data Security. Provider shall implement and maintain administrative, technical, and physical \
safeguards consistent with industry standards (SOC 2 Type II) to protect Client data.

ARTICLE 5: INTELLECTUAL PROPERTY

5.1 Client IP. All intellectual property provided by Client to Provider ("Client IP") shall \
remain the sole property of Client.

5.2 Work Product. All deliverables and work product created by Provider specifically for Client \
under a SOW ("Work Product") shall be owned by Client upon full payment. Provider hereby assigns \
all right, title, and interest in such Work Product to Client.

5.3 Provider IP. Provider retains all rights in its pre-existing intellectual property, tools, \
methodologies, and general know-how ("Provider IP"). Provider grants Client a non-exclusive, \
perpetual, irrevocable license to use any Provider IP incorporated into the Work Product.

5.4 Third-Party Components. Provider shall identify all third-party or open-source components \
incorporated into deliverables and ensure appropriate licenses are obtained.

ARTICLE 6: LIMITATION OF LIABILITY

6.1 Cap. EXCEPT FOR BREACHES OF ARTICLE 3 (CONFIDENTIALITY) AND ARTICLE 4 (DATA PROTECTION), \
NEITHER PARTY'S TOTAL LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE \
BY CLIENT IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM.

6.2 Enhanced Cap for Data Breaches. For breaches of Article 4 (Data Protection), Provider's \
liability shall not exceed two (2) times the annual fees paid or payable under this Agreement.

6.3 Exclusion of Consequential Damages. NEITHER PARTY SHALL BE LIABLE FOR ANY INDIRECT, \
INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOST PROFITS, REGARDLESS OF \
THE FORM OF ACTION, EXCEPT FOR (A) BREACHES OF CONFIDENTIALITY, (B) WILLFUL MISCONDUCT, OR \
(C) INDEMNIFICATION OBLIGATIONS.

ARTICLE 7: INDEMNIFICATION

7.1 Provider Indemnification. Provider shall defend, indemnify, and hold harmless Client from \
any third-party claims arising from: (a) Provider's breach of this Agreement; (b) Provider's \
negligence or willful misconduct; (c) infringement of third-party intellectual property rights \
by the Services or Work Product; (d) Provider's violation of applicable laws.

7.2 Client Indemnification. Client shall defend, indemnify, and hold harmless Provider from \
third-party claims arising from Client's use of the Services in violation of this Agreement or \
applicable law.

7.3 Indemnification Procedures. The indemnified party shall: (a) provide prompt written notice; \
(b) grant the indemnifying party sole control of the defense; (c) provide reasonable cooperation.

ARTICLE 8: INSURANCE

8.1 Required Coverage. Provider shall maintain: (a) commercial general liability insurance of at \
least $2,000,000 per occurrence; (b) professional liability / errors and omissions insurance of \
at least $5,000,000; (c) cyber liability insurance of at least $5,000,000; (d) workers' \
compensation as required by law.

8.2 Evidence. Provider shall provide certificates of insurance upon request.

ARTICLE 9: FORCE MAJEURE

9.1 Definition. Neither party shall be liable for failure to perform its obligations (other than \
payment obligations) due to events beyond its reasonable control, including natural disasters, \
pandemics, war, terrorism, government actions, or infrastructure failures ("Force Majeure Event").

9.2 Notice. The affected party shall provide written notice within five (5) business days of the \
Force Majeure Event, describing the event and expected duration.

9.3 Mitigation. The affected party shall use commercially reasonable efforts to mitigate the \
impact and resume performance as soon as practicable.

9.4 Extended Force Majeure. If a Force Majeure Event continues for more than sixty (60) days, \
either party may terminate this Agreement upon fifteen (15) days written notice.

ARTICLE 10: GENERAL PROVISIONS

10.1 Governing Law. This Agreement shall be governed by the laws of the State of New York.

10.2 Dispute Resolution. Any dispute shall first be submitted to mediation. If mediation fails \
within sixty (60) days, either party may pursue litigation in the state or federal courts located \
in New York County, New York.

10.3 Notices. All notices shall be in writing and delivered by certified mail, overnight courier, \
or email to the addresses specified in the applicable SOW.

10.4 Assignment. Neither party may assign this Agreement without the prior written consent of \
the other party, except in connection with a merger, acquisition, or sale of substantially all \
of its assets.

10.5 Entire Agreement. This Agreement, together with all SOWs, constitutes the entire agreement \
between the parties.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.

Acme Technology Solutions, Inc.        GlobalCorp Industries, LLC
By: _________________________         By: _________________________
Name: Jane Smith                       Name: Robert Chen
Title: CEO                             Title: General Counsel
Date: January 15, 2026                 Date: January 15, 2026
"""


def create_sample_docx() -> None:
    """Create a synthetic SaaS Agreement as DOCX."""
    import docx

    doc = docx.Document()
    for para_text in MSA_TEXT.strip().split("\n\n"):
        text = para_text.strip()
        if not text:
            continue
        if text.startswith("ARTICLE") or text == "MASTER SERVICES AGREEMENT":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        else:
            doc.add_paragraph(text)

    output_path = SAMPLE_DIR / "sample_msa.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_sample_pdf() -> None:
    """Create a synthetic MSA as PDF using reportlab if available, otherwise a text-based PDF."""
    output_path = SAMPLE_DIR / "sample_msa.pdf"

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for para_text in MSA_TEXT.strip().split("\n\n"):
            text = para_text.strip()
            if not text:
                continue
            if text.startswith("ARTICLE") or text == "MASTER SERVICES AGREEMENT":
                story.append(Paragraph(text, styles["Heading2"]))
            else:
                story.append(Paragraph(text, styles["Normal"]))
            story.append(Spacer(1, 6))

        doc.build(story)
    except ImportError:
        # Fallback: create a minimal valid PDF with the text
        _create_simple_pdf(output_path, MSA_TEXT)

    print(f"Created: {output_path}")


def _create_simple_pdf(output_path: Path, text: str) -> None:
    """Create a minimal PDF file with text content."""
    # Minimal PDF structure
    lines = text.encode("latin-1", errors="replace")
    content = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R"
        b"/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Courier>>endobj\n"
    )

    # Wrap text for PDF content stream
    wrapped_lines = []
    for line in text.split("\n"):
        while len(line) > 80:
            wrapped_lines.append(line[:80])
            line = line[80:]
        wrapped_lines.append(line)

    # Only include first ~60 lines to fit on one page
    display_lines = wrapped_lines[:60]
    stream_content = b"BT\n/F1 8 Tf\n"
    y = 750
    for line in display_lines:
        safe_line = line.replace("(", "\\(").replace(")", "\\)")
        stream_content += f"1 0 0 1 36 {y} Tm ({safe_line}) Tj\n".encode("latin-1", errors="replace")
        y -= 10

    stream_content += b"ET\n"

    stream_obj = (
        f"4 0 obj<</Length {len(stream_content)}>>stream\n".encode()
        + stream_content
        + b"endstream endobj\n"
    )

    full = content + stream_obj
    xref_pos = len(full)
    full += (
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
    )
    full += f"{len(content):010d} 00000 n \n".encode()
    full += b"0000000000 00000 n \n"  # placeholder for font
    full += b"trailer<</Size 6/Root 1 0 R>>\n"
    full += f"startxref\n{xref_pos}\n%%EOF\n".encode()

    output_path.write_bytes(full)


if __name__ == "__main__":
    os.chdir(Path(__file__).parent)
    create_sample_docx()
    create_sample_pdf()
    print("Sample contracts created successfully.")
