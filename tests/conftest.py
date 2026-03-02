import io
import os
from pathlib import Path
from unittest.mock import patch

import pytest
from fastapi.testclient import TestClient

from contract_simulator.core.config import Settings


@pytest.fixture
def settings():
    """Create test settings with a dummy API key."""
    scenarios_dir = str(Path(__file__).parent.parent / "data" / "scenarios")
    with patch.dict(os.environ, {"ANTHROPIC_API_KEY": "test-key-not-real"}):
        return Settings(
            anthropic_api_key="test-key-not-real",
            scenarios_dir=scenarios_dir,
        )


@pytest.fixture
def test_client(settings):
    """Create a FastAPI test client with overridden settings."""
    from contract_simulator.api.main import app
    from contract_simulator.core.config import get_settings

    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)
    yield client
    app.dependency_overrides.clear()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Create minimal valid PDF bytes for testing."""
    content = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R"
        b"/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Courier>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 12 Tf 100 700 Td (Test contract) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 6\n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n0\n%%EOF\n"
    )
    return content


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Create minimal valid DOCX bytes for testing."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("MASTER SERVICES AGREEMENT")
    doc.add_paragraph("This Agreement is between Party A and Party B.")
    doc.add_paragraph("1. CONFIDENTIALITY")
    doc.add_paragraph("Each party agrees to maintain confidentiality.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_parsed_contract() -> dict:
    """Return a sample parsed contract dict for simulation tests."""
    return {
        "contract_title": "Master Services Agreement",
        "parties": ["Acme Technology Solutions, Inc.", "GlobalCorp Industries, LLC"],
        "effective_date": "January 15, 2026",
        "raw_text": "Sample contract text...",
        "clauses": [
            {
                "id": "clause_1",
                "title": "Confidentiality",
                "section_number": "3.1",
                "content": "Each party agrees to maintain confidentiality of the other party's information.",
                "clause_type": "confidentiality",
            },
            {
                "id": "clause_2",
                "title": "Data Breach Notification",
                "section_number": "4.2",
                "content": "Provider shall notify Client within 48 hours of discovery of a data breach.",
                "clause_type": "data_protection",
            },
            {
                "id": "clause_3",
                "title": "Limitation of Liability",
                "section_number": "6.1",
                "content": "Neither party's total liability shall exceed the fees paid in the 12 months preceding the claim.",
                "clause_type": "liability",
            },
        ],
    }
